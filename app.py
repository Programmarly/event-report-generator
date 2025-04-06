import os
import requests
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from flask import Flask, request, jsonify, send_file
import os
import easyocr
import pandas as pd
from transformers import pipeline
from sklearn.feature_extraction.text import CountVectorizer
from collections import Counter
from ollama import chat
from ollama import ChatResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import markdown2
from docx import Document
import warnings
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize OCR reader
easyocr_reader = easyocr.Reader(['en'])

# Sentiment Analysis Pipeline
sentiment_analyzer = pipeline("sentiment-analysis")

def get_user_folder(email):
    """Creates and returns a user-specific directory for file storage."""
    print("fetching email", email)
    folder = os.path.join(UPLOAD_FOLDER, email.replace("@", "_").replace(".", "_"))
    os.makedirs(folder, exist_ok=True)
    return folder

def extract_text(image_path):
    """Extracts text from an image using OCR."""
    print("Extracting text from image...")
    result = easyocr_reader.readtext(image_path, detail=0)
    return " ".join(result)

def process_text(text):
    """Fixes grammar and enhances text using AI."""
    print("Processing text...")
    response: ChatResponse = chat(model='gemma2:9b', messages=[
        {'role': 'user',
    'content': f"following is my raw text after OCR, fix any grammatical errors and polish it to look more professional for a brochure invite. Just provide me with the polished text. {text}" },

    ])
    return response.message.content

def analyze_feedback(feedback):
  print("Analyzing feedback...")
  feedback_data = pd.read_csv(feedback)
  average_rating = feedback_data["Rating"].mean()
  rating_distribution = feedback_data["Rating"].value_counts(normalize=True) * 100
  
  sentiment_analyzer = pipeline("sentiment-analysis", model="distilbert-base-uncased-finetuned-sst-2-english")
  feedback_data["Sentiment"] = feedback_data["Feedback"].apply(lambda x: sentiment_analyzer(x)[0]["label"])
  sentiment_distribution = feedback_data["Sentiment"].value_counts(normalize=True) * 100
  
  vectorizer = CountVectorizer(stop_words="english")
  X = vectorizer.fit_transform(feedback_data["Feedback"])
  word_counts = Counter(dict(zip(vectorizer.get_feature_names_out(), X.toarray().sum(axis=0))))
  common_themes = word_counts.most_common(5)
  quant_feedback = f"""
  1. Total Attendees: {len(feedback_data)}
2. Average Rating: {average_rating:.2f}
3. Rating Distribution:
   - 5 stars: {rating_distribution.get(5, 0):.1f}%
   - 4 stars: {rating_distribution.get(4, 0):.1f}%
   - 3 stars: {rating_distribution.get(3, 0):.1f}%
4. Sentiment Analysis:
   - Positive feedback: {sentiment_distribution.get('POSITIVE', 0):.1f}%
   - Neutral feedback: {sentiment_distribution.get('NEUTRAL', 0):.1f}%
   - Negative feedback: {sentiment_distribution.get('NEGATIVE', 0):.1f}%
5. Common Themes:
   - {', '.join([theme[0] for theme in common_themes])}
   """ 
   
  prompt = f"""
Provided is  the following feedback insights professionally for an event report:
  {quant_feedback}

Analyse {feedback_data} and the above insights and provide a detailed summary of the feedback, highlighting prominent and critical audience comments.
Do numerical analysis as well, highlighting how many attendees rated the event positively, neutrally, and negatively. AND WHAT IS THE RATING DISTRIBUTION AS WELL.
"""
  response: ChatResponse = chat(model='gemma2:9b', messages=[
  {
    'role': 'user',
    'content': prompt },
  ])


  return quant_feedback, response.message.content

def fetch_linkedin_profile(linkedin_profile_url,api_key='kEJCzBgzdn8qShvstBTpgA'):
    """
    Fetches LinkedIn profile data and returns the concatenated details as a string.

    Args:
        api_key (str): API key for Proxycurl.
        linkedin_profile_url (str): LinkedIn profile URL.

    Returns:
        str: Concatenated full name, occupation, headline, and summary separated by newlines.
    """
    api_endpoint = 'https://nubela.co/proxycurl/api/v2/linkedin'
    headers = {'Authorization': 'Bearer ' + api_key}

    try:
        response = requests.get(api_endpoint,
                                params={'url': linkedin_profile_url, 'skills': 'include'},
                                headers=headers)
        response.raise_for_status()
        profile_data = response.json()

        full_name = profile_data.get('full_name', 'N/A')
        occupation = profile_data.get('occupation', 'N/A')
        headline = profile_data.get('headline', 'N/A')
        summary = profile_data.get('summary', 'N/A')

        return f"Speaker's name is: {full_name}. Headline is {headline}.His occupation is: {occupation}.Professional Summary of the Speaker: {summary}"

    except requests.exceptions.RequestException as e:
        return f"Error: {e}" 

def generate_report(polished_text,quant_data, linkedin_profile, feedback_summary):
  print("Generating report...")
  prompt = f"""You are an AI assistant tasked with creating a professional and detailed **After-Event Report**. Below are the inputs you have:

### Inputs:
1. **Event Details**: 
   {polished_text}

2. **Quantitative Attendance Analysis**: 
   {quant_data}

3. **Feedback Summary**:
   {feedback_summary}

4. **LinkedIn Profile Summary**:
   {linkedin_profile} #In your own words, summarize the speaker details and their contributions to the event / or in life in general.
   
### Task:
Using the provided details, generate a **comprehensive After-Event Report** in the following structure:

---

## **After-Event Report Template**

### 1. Title Page
- **Event Name**: Provide the official event name.
- **Event Date**: Mention the event date.
- **Event Location**: Specify the venue/location.
- **Organized by**: Mention the organizing entity or entities.

---

### 2. Executive Summary
- Write a concise paragraph summarizing the purpose of the event, its significance, and the overall outcome.
- Highlight key numbers (e.g., number of attendees, average rating) and major achievements.
- Include a brief mention of the LinkedIn profile summary, if relevant.
- Add  {linkedin_profile} to the summary.
- Provide a link to the LinkedIn profile for further details.

---

### 3. Event Objectives
- List the primary objectives or goals of the event in a bulleted format.
- Mention why the event was significant and the intended outcomes.

---

### 4. Event Details
- Provide the date, time, and location of the event which you received from .
- Mention the key agenda highlights - if any. (e.g., workshop sessions, keynote speeches, networking opportunities).

---

### 6. Attendee Statistics (gather data from {quant_data})
- Total number of attendees.
- Satisfaction ratings (average rating and percentage breakdown of 5 stars, 4 stars, etc.).

---

### 7. Key Takeaways
- Summarize the key insights, learnings, or outcomes participants gained from the event.

---

### 8. Feedback Summary
- Include the quantitative analysis from feedback:
   - Average rating.
   - Sentiment distribution (positive, neutral, negative).
- Highlight qualitative insights:
   - Common themes or recurring feedback points.
   - Notable participant quotes or testimonials.

---

### 10. Conclusion
- Provide a final summary of the event's success and impact.
- Acknowledge contributors, sponsors, or key supporters.

---

### 11. Recommendations for Future Events
- Include actionable suggestions based on feedback and event observations.

---

### Additional Guidelines:
1. Use a **formal and professional tone** throughout the report.
2. Organize the content logically, ensuring it flows naturally between sections.
3. Incorporate any quantitative data (from feedback_summary) into tables or bulleted points for clarity.
4. Highlight exceptional feedback or key moments from the event that would leave a lasting impression on the reader.
5. Ensure every section is detailed and supports the overall narrative of the event's success.

---"""
  response: ChatResponse = chat(model='gemma2:9b', messages=[
  {
    'role': 'user',
    'content': prompt },
  ])
  print("Generating report")
  return response.message.content

def add_images_to_docx(doc, image_paths, images_per_row=3, image_width_inches=3.13):
    """
    Adds a photo gallery to the DOCX document in a grid format.
    Starts on a new page with heading 'Photo Gallery'.
    Each page can contain 2 rows and 3 columns (6 images per page).
    """
    if not image_paths:
        return doc

    # Add page break and heading
    doc.add_page_break()
    heading = doc.add_paragraph("Photo Gallery")
    heading.style = 'Heading 1'
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = None
    count = 0

    for i, image_path in enumerate(image_paths):
        if not os.path.exists(image_path):
            continue

        if count % 6 == 0:  # New page every 6 images (2 rows x 3 cols)
            if table:  # If table already exists, add page break before next table
                doc.add_page_break()
                heading = doc.add_paragraph("Photo Gallery (contd.)")
                heading.style = 'Heading 1'
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = doc.add_table(rows=2, cols=3)
            table.autofit = False
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(image_width_inches)

        row = (count % 6) // 3
        col = (count % 6) % 3
        cell = table.cell(row, col)

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(image_width_inches))

        count += 1

    return doc

def insert_text_into_docx(template_path, output_path, image_path, report_text, image_paths=None):
    """
    Inserts report text into a DOCX template while converting Markdown to readable Word formatting.
    """
    # Open the Word document template
    doc = Document(template_path)
    
    # Convert Markdown to HTML and parse it into plain text with formatting
    html = markdown2.markdown(report_text)
    
    # Split the Markdown content into lines for processing
    lines = report_text.split("\n")

    # Find and replace the placeholder "{{REPORT_CONTENT}}" with formatted content
    for para in doc.paragraphs:
        if "{{REPORT_CONTENT}}" in para.text:
            # Remove the placeholder text
            para.text = ""
            
            # Process each line of the Markdown content
            for line in lines:
                if line.startswith("# "):  # H1 Heading
                    heading = doc.add_paragraph(line[2:])
                    heading.style = 'Heading 1'
                elif line.startswith("## "):  # H2 Heading
                    heading = doc.add_paragraph(line[3:])
                    heading.style = 'Heading 2'
                elif line.startswith("### "):  # H3 Heading
                    heading = doc.add_paragraph(line[4:])
                    heading.style = 'Heading 3'
                elif line.startswith("- "):  # Bullet points
                    bullet = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith("** "):  # Numbered list (basic example)
                    bold = doc.add_paragraph(line[3:], style='bold')
                elif line.strip() == "":  # Blank line (add spacing)
                    doc.add_paragraph("")
                else:  # Normal paragraph text
                    normal_text = doc.add_paragraph(line)
                    normal_text.style.font.size = Pt(12)  # Set font size for normal text
            
            break  # Exit after replacing the placeholder
    if image_paths:
      doc = add_images_to_docx(doc, image_path)
    # Save the modified document
    doc.save(output_path)
    return output_path

def add_images_to_docx(doc, image_paths):
    """
    Adds images to the DOCX document.
    """
    for image_path in image_paths:
        if os.path.exists(image_path):
            doc.add_picture(image_path)
            doc.add_paragraph("")  # Add a blank paragraph for spacing
    return doc
  
@app.route('/generate_report', methods=['POST'])
def generate_report_api():
    """Main endpoint to generate the report in a single request."""
    email = request.form.get('email')  
    if not email:
        return jsonify({"error": "Email is required"}), 400
    user_folder = get_user_folder(email)
    event_image = request.files.get('brochure_image')
    feedback_file = request.files.get('feedback')
    linkedin_url = request.form.get('linkedin_url')
    images = request.files.getlist('images')
    image_paths = []

    if images:
      img_dir = os.path.join(f"uploads/{user_folder}/images")
      os.makedirs(img_dir, exist_ok=True)
      for idx, img in enumerate(images):
          img_path = os.path.join(img_dir, f"image_{idx}.jpg")
          img.save(img_path)
          image_paths.append(img_path)
          
        
    template_file = request.files.get('templateFile')
    if not event_image or not feedback_file:
        return jsonify({"error": "Both image and feedback file are required"}), 400


    #Save files
    event_image_path = os.path.join(user_folder, "image.jpg")
    event_image.save(event_image_path)
    
    feedback_path = os.path.join(user_folder, "feedback.csv")
    feedback_file.save(feedback_path)
    
    if template_file:
        template_path = os.path.join(user_folder, "template.docx")
        template_file.save(template_path)
    else:
        template_path = "raw_template.docx"
        with open(template_path, 'w') as f:
            f.write("{{REPORT_CONTENT}}")
    

    # Step 1: Extract text from image
    extracted_text = extract_text(event_image_path)
    
    print("Finished extracting text")

    # Step 2: Process the extracted text
    polished_text = process_text(extracted_text)

    print("Finished processing text")
    # Step 3: Analyze feedback data
    quant_data, feedback_summary = analyze_feedback(feedback_path)
    
    print("Finished analyzing feedback")
    # Step 3: Fetch LinkedIn profile details
    speaker_profile_details = fetch_linkedin_profile(linkedin_url)
    print("Finished fetching LinkedIn profile details")
    # Step 4: Generate the final report
    report_text = generate_report(polished_text,quant_data,speaker_profile_details, feedback_summary)
    
    print("Finished generating report")
      
    # Step 6: Generate and save the DOCX
    docx_path = insert_text_into_docx(os.path.join(user_folder,"template.docx"), os.path.join(user_folder, "event_report.docx"), image_paths, report_text)
    
    print("Finished generating DOCX report")

    # Step 6: Return the PDF file
    return send_file(docx_path, as_attachment=True, download_name="event_report.docx")


if __name__ == '__main__':
    print("Starting Flask server...")
    app.run(port=8000,debug=True)
